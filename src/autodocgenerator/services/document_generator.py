from __future__ import annotations

import re
from copy import deepcopy
from dataclasses import dataclass
from datetime import date, datetime
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from PIL import Image, ImageOps, UnidentifiedImageError

from autodocgenerator.domain.exceptions import DocumentGenerationError
from autodocgenerator.domain.models import SourceImage

INDONESIAN_MONTH_NAMES = {
    1: "JANUARI", 2: "FEBRUARI", 3: "MARET", 4: "APRIL",
    5: "MEI", 6: "JUNI", 7: "JULI", 8: "AGUSTUS",
    9: "SEPTEMBER", 10: "OKTOBER", 11: "NOVEMBER", 12: "DESEMBER",
}


@dataclass(slots=True, frozen=True)
class DocumentGenerationSettings:
    company_name: str
    bank_name: str = "BCA"
    title_font: str = "Times New Roman"
    title_font_size_pt: float = 12.0
    title_bold: bool = False
    first_page_image_width_cm: float = 13.70
    first_page_image_height_cm: float = 13.70
    other_page_image_width_cm: float = 14.16
    other_page_image_right_ruler_cm: float = 17.00
    other_page_image_height_cm: float = 14.16
    first_page_first_image_top_cm: float = 1.40
    first_page_second_image_top_cm: float = 15.25
    other_page_first_image_top_cm: float = 0.65
    other_page_second_image_top_cm: float = 14.90
    page_width_cm: float = 21.0
    page_height_cm: float = 29.7
    page_margin_top_cm: float = 0.75
    page_margin_bottom_cm: float = 2.54
    page_margin_left_cm: float = 2.54
    page_margin_right_cm: float = 2.54
    transfer_canvas_size_px: int = 1800
    receipt_max_width_cm: float = 14.0
    receipt_max_height_cm: float = 22.0
    pdf_first_page_max_height_cm: float = 20.8
    pdf_title_font_size_pt: float = 11.0
    pdf_title_bold: bool = True
    pdf_title_space_after_pt: float = 6.0
    jpeg_quality: int = 95
    picture_border_width_pt: float = 1.5
    picture_border_color_hex: str = "000000"

    def __post_init__(self) -> None:
        if not self.company_name.strip():
            raise ValueError("company_name cannot be empty.")

        if self.picture_border_width_pt < 0:
            raise ValueError(
                "picture_border_width_pt must be greater than or equal to zero."
            )

        color = self.picture_border_color_hex.strip().lstrip("#")
        if not re.fullmatch(r"[0-9A-Fa-f]{6}", color):
            raise ValueError(
                "picture_border_color_hex must contain exactly six hex digits."
            )

        right_edge_from_page_cm = (
            self.page_margin_left_cm
            + self.other_page_image_right_ruler_cm
        )
        calculated_left_cm = (
            right_edge_from_page_cm
            - self.other_page_image_width_cm
        )

        if calculated_left_cm < 0:
            raise ValueError(
                "The page-two image target places the image "
                "outside the left edge of the page."
            )

        if right_edge_from_page_cm > self.page_width_cm:
            raise ValueError(
                "other_page_image_right_ruler_cm places the image "
                "outside the right edge of the page."
            )


class DocumentGenerator:
    """Generate two transfer proofs per page and append receipts last."""

    def __init__(self, *, settings: DocumentGenerationSettings) -> None:
        self._settings = settings

    def generate(
        self,
        source_images: list[SourceImage],
        output_directory: Path,
        *,
        document_date: date | datetime | None = None,
    ) -> Path:
        usable = [
            image
            for image in source_images
            if image.processed_path is not None and image.processed_path.exists()
        ]
        if not usable:
            raise DocumentGenerationError("Tidak ada gambar hasil proses yang valid.")

        resolved_date = self._resolve_date(usable, document_date)
        transfers = [image for image in usable if not image.is_real_receipt]
        receipts = [image for image in usable if image.is_real_receipt]

        document = Document()
        self._configure(document)
        self._add_title(document, self._build_title(resolved_date))
        self._add_transfer_images(document, transfers)
        self._add_receipts(document, receipts, has_prior_content=bool(transfers))

        document_directory = output_directory.expanduser().resolve() / "document"
        document_directory.mkdir(parents=True, exist_ok=True)
        output_path = document_directory / self._filename(resolved_date)
        try:
            document.save(output_path)
        except OSError as error:
            raise DocumentGenerationError(
                f"Gagal menyimpan dokumen Word: {output_path}"
            ) from error
        return output_path

    def _configure(self, document: DocumentObject) -> None:
        section = document.sections[0]
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Cm(self._settings.page_width_cm)
        section.page_height = Cm(self._settings.page_height_cm)
        section.top_margin = Cm(self._settings.page_margin_top_cm)
        section.bottom_margin = Cm(self._settings.page_margin_bottom_cm)
        section.left_margin = Cm(self._settings.page_margin_left_cm)
        section.right_margin = Cm(self._settings.page_margin_right_cm)
        style = document.styles["Normal"]
        style.font.name = self._settings.title_font
        style.font.size = Pt(self._settings.title_font_size_pt)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), self._settings.title_font)

    def _add_title(self, document: DocumentObject, title: str) -> None:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(title)
        run.bold = self._settings.title_bold
        run.font.name = self._settings.title_font
        run.font.size = Pt(self._settings.title_font_size_pt)
        run._element.rPr.rFonts.set(qn("w:eastAsia"), self._settings.title_font)

    def _add_transfer_images(
        self,
        document: DocumentObject,
        images: list[SourceImage],
    ) -> None:
        for index, source_image in enumerate(images):
            page_index = index // 2
            position = index % 2
            first_page = page_index == 0

            if index == 0:
                paragraph = document.add_paragraph()
            elif position == 0:
                document.add_page_break()
                paragraph = document.add_paragraph()
            else:
                paragraph = document.paragraphs[-1]

            if first_page:
                width = self._settings.first_page_image_width_cm
                height = self._settings.first_page_image_height_cm
                x = max(0.0, (self._settings.page_width_cm - width) / 2)
                y = (
                    self._settings.first_page_first_image_top_cm
                    if position == 0
                    else self._settings.first_page_second_image_top_cm
                )
            else:
                width = self._settings.other_page_image_width_cm
                height = self._settings.other_page_image_height_cm
                x = (
                    self._settings.page_margin_left_cm
                    + self._settings.other_page_image_right_ruler_cm
                    - width
                )
                y = (
                    self._settings.other_page_first_image_top_cm
                    if position == 0
                    else self._settings.other_page_second_image_top_cm
                )

            stream = self._transfer_stream(source_image)
            self._add_floating_picture(
                paragraph=paragraph,
                image_stream=stream,
                width_cm=width,
                height_cm=height,
                horizontal_position_cm=x,
                vertical_position_cm=y,
                picture_id=index + 1,
                picture_name=source_image.filename,
            )

    def _add_receipts(
        self,
        document: DocumentObject,
        receipts: list[SourceImage],
        *,
        has_prior_content: bool,
    ) -> None:
        for index, source_image in enumerate(receipts):
            if has_prior_content or index > 0:
                document.add_page_break()

            add_pdf_title = source_image.is_first_pdf_page

            if add_pdf_title:
                self._add_pdf_title(
                    document,
                    source_image.resolved_pdf_title,
                )

            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)

            image_path = source_image.processed_path

            if image_path is None:
                continue

            maximum_height = (
                self._settings.pdf_first_page_max_height_cm
                if add_pdf_title
                else self._settings.receipt_max_height_cm
            )
            width_cm, height_cm = self._fit_dimensions(
                image_path,
                self._settings.receipt_max_width_cm,
                maximum_height,
            )
            run = paragraph.add_run()
            inline_shape = run.add_picture(
                str(image_path),
                width=Cm(width_cm),
                height=Cm(height_cm),
            )
            self._apply_picture_border(
                inline_shape._inline.graphic
            )

    def _add_pdf_title(
        self,
        document: DocumentObject,
        title: str,
    ) -> None:
        """Add the PDF filename once above the first rendered PDF page."""
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(
            self._settings.pdf_title_space_after_pt
        )
        paragraph.paragraph_format.keep_with_next = True

        run = paragraph.add_run(title)
        run.bold = self._settings.pdf_title_bold
        run.font.name = self._settings.title_font
        run.font.size = Pt(
            self._settings.pdf_title_font_size_pt
        )
        run._element.rPr.rFonts.set(
            qn("w:eastAsia"),
            self._settings.title_font,
        )

    def _transfer_stream(self, source_image: SourceImage) -> BytesIO:
        image_path = source_image.processed_path
        if image_path is None:
            raise DocumentGenerationError(
                f"Processed path tidak tersedia untuk {source_image.filename}."
            )
        try:
            with Image.open(image_path) as source:
                prepared = ImageOps.exif_transpose(source).convert("RGB")
                size = self._settings.transfer_canvas_size_px
                full_frame = prepared.resize((size, size), Image.Resampling.LANCZOS)
                stream = BytesIO()
                full_frame.save(
                    stream,
                    format="JPEG",
                    quality=self._settings.jpeg_quality,
                    optimize=True,
                    dpi=(96, 96),
                )
                stream.seek(0)
                return stream
        except (UnidentifiedImageError, OSError) as error:
            raise DocumentGenerationError(
                f"Gambar hasil proses tidak dapat dibuka: {image_path}"
            ) from error

    @staticmethod
    def _fit_dimensions(
        path: Path,
        max_width: float,
        max_height: float,
    ) -> tuple[float, float]:
        """Preserve the original physical size and aspect ratio.

        REAL RECEIPT files are never stretched to fill the page. They are
        only reduced proportionally when their natural size is larger than
        the usable Word page area.
        """
        try:
            with Image.open(path) as image:
                width_px, height_px = image.size
                dpi_value = image.info.get("dpi", (96.0, 96.0))
        except (UnidentifiedImageError, OSError) as error:
            raise DocumentGenerationError(
                f"Gambar tidak valid: {path}"
            ) from error

        try:
            dpi_x = float(dpi_value[0])
            dpi_y = float(dpi_value[1])
        except (IndexError, TypeError, ValueError):
            dpi_x = 96.0
            dpi_y = 96.0

        if not 1 <= dpi_x <= 1_200:
            dpi_x = 96.0

        if not 1 <= dpi_y <= 1_200:
            dpi_y = 96.0

        natural_width_cm = width_px / dpi_x * 2.54
        natural_height_cm = height_px / dpi_y * 2.54
        scale = min(
            max_width / natural_width_cm,
            max_height / natural_height_cm,
            1.0,
        )

        return (
            natural_width_cm * scale,
            natural_height_cm * scale,
        )

    def _add_floating_picture(
        self,
        *,
        paragraph: Any,
        image_stream: BytesIO,
        width_cm: float,
        height_cm: float,
        horizontal_position_cm: float,
        vertical_position_cm: float,
        picture_id: int,
        picture_name: str,
    ) -> None:
        run = paragraph.add_run()
        inline_shape = run.add_picture(
            image_stream,
            width=Cm(width_cm),
            height=Cm(height_cm),
        )
        inline = inline_shape._inline
        graphic = deepcopy(inline.graphic)
        self._apply_picture_border(graphic)
        anchor = OxmlElement("wp:anchor")
        for key, value in {
            "distT": "0", "distB": "0", "distL": "0", "distR": "0",
            "simplePos": "0", "relativeHeight": str(251658240 + picture_id),
            "behindDoc": "0", "locked": "0", "layoutInCell": "1",
            "allowOverlap": "1",
        }.items():
            anchor.set(key, value)
        simple = OxmlElement("wp:simplePos")
        simple.set("x", "0")
        simple.set("y", "0")
        anchor.append(simple)

        horizontal = OxmlElement("wp:positionH")
        horizontal.set("relativeFrom", "page")
        horizontal_offset = OxmlElement("wp:posOffset")
        horizontal_offset.text = str(int(Cm(horizontal_position_cm)))
        horizontal.append(horizontal_offset)
        anchor.append(horizontal)

        vertical = OxmlElement("wp:positionV")
        vertical.set("relativeFrom", "page")
        vertical_offset = OxmlElement("wp:posOffset")
        vertical_offset.text = str(int(Cm(vertical_position_cm)))
        vertical.append(vertical_offset)
        anchor.append(vertical)

        extent = OxmlElement("wp:extent")
        extent.set("cx", str(int(Cm(width_cm))))
        extent.set("cy", str(int(Cm(height_cm))))
        anchor.append(extent)

        effect = OxmlElement("wp:effectExtent")
        for key in ("l", "t", "r", "b"):
            effect.set(key, "0")
        anchor.append(effect)
        anchor.append(OxmlElement("wp:wrapNone"))

        document_properties = OxmlElement("wp:docPr")
        document_properties.set("id", str(picture_id))
        document_properties.set("name", self._safe_name(picture_name))
        anchor.append(document_properties)

        frame_properties = OxmlElement("wp:cNvGraphicFramePr")
        locks = OxmlElement("a:graphicFrameLocks")
        locks.set("noChangeAspect", "0")
        frame_properties.append(locks)
        anchor.append(frame_properties)
        anchor.append(graphic)
        inline.getparent().replace(inline, anchor)

    def _apply_picture_border(
        self,
        graphic: Any,
    ) -> None:
        """Apply a true Word picture outline to an inserted image."""
        width_pt = self._settings.picture_border_width_pt

        if width_pt <= 0:
            return

        shape_properties = graphic.xpath(
            ".//pic:spPr"
        )

        if not shape_properties:
            raise DocumentGenerationError(
                "Properti gambar Word tidak ditemukan saat menambahkan border."
            )

        shape_property = shape_properties[0]

        for existing_line in shape_property.xpath(
            "./a:ln"
        ):
            shape_property.remove(existing_line)

        line = OxmlElement("a:ln")
        line.set(
            "w",
            str(round(width_pt * 12_700)),
        )

        solid_fill = OxmlElement("a:solidFill")
        color = OxmlElement("a:srgbClr")
        color.set(
            "val",
            self._settings.picture_border_color_hex
            .strip()
            .lstrip("#")
            .upper(),
        )
        solid_fill.append(color)
        line.append(solid_fill)

        dash = OxmlElement("a:prstDash")
        dash.set("val", "solid")
        line.append(dash)

        shape_property.append(line)

    @staticmethod
    def _resolve_date(
        images: list[SourceImage],
        explicit: date | datetime | None,
    ) -> date:
        if explicit is not None:
            return explicit.date() if isinstance(explicit, datetime) else explicit
        datetimes = [
            image.transaction_datetime
            for image in images
            if image.transaction_datetime is not None
        ]
        return min(datetimes).date() if datetimes else date.today()

    def _build_title(self, document_date: date) -> str:
        return (
            f"BUKTI PENGELUARAN TGL {document_date.day} "
            f"{INDONESIAN_MONTH_NAMES[document_date.month]} {document_date.year} - "
            f"{self._settings.company_name.strip()}\n"
            f"({self._settings.bank_name.strip()})"
        )

    def _filename(self, document_date: date) -> str:
        raw = (
            f"BUKTI PENGELUARAN TGL {document_date.day} "
            f"{INDONESIAN_MONTH_NAMES[document_date.month]} {document_date.year}.docx"
        )
        return re.sub(r'[<>:"/\\|?*]+', "_", raw)

    @staticmethod
    def _safe_name(name: str) -> str:
        return re.sub(r"[^A-Za-z0-9_. -]+", "_", name)[:120] or "image"
